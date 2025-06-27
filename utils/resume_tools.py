from utils.gemini_app import gemini_chat
from docx import Document
from docx.shared import Inches
import uuid
import os

def generate_resume(user_input, uploaded_image=None):
    prompt_template = open("prompts/resume_generator.txt", "r").read()
    filled_prompt = prompt_template.replace("{{user_input}}", user_input)
    resume_text = gemini_chat(filled_prompt, system="You are a professional resume writer.")

    doc = Document()
    doc.add_heading('Resume', 0)

    if uploaded_image:
        doc.add_picture(uploaded_image, width=Inches(1.5))

    for line in resume_text.split("\n"):
        if line.strip() == "":
            continue
        if line.endswith(":"):
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    filename = f"generated_resume_{uuid.uuid4().hex}"
    os.makedirs("generated", exist_ok=True)
    docx_path = os.path.join("generated", f"{filename}.docx")
    html_path = os.path.join("generated", f"{filename}.html")

    doc.save(docx_path)

    html_content = f"""
    <html>
    <head><meta charset="utf-8"><title>Resume</title></head>
    <body style="font-family:Arial;padding:40px;line-height:1.6;color:#333;">
    <h1 style="text-align:center;color:#2E86C1;">{user_input.splitlines()[0]}</h1>
    {''.join(f'<h2>{line}</h2>' if line.endswith(':') else f'<p>{line}</p>' for line in resume_text.splitlines() if line.strip())}
    </body></html>
    """
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    return resume_text, docx_path, None, html_path
